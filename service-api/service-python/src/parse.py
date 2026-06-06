import io
import re
from dataclasses import dataclass, field
from typing import Any

import docx
import fitz
import requests
from bs4 import BeautifulSoup


@dataclass
class DocumentBlock:
    block_type: str
    raw_text: str
    clean_text: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    raw_text: str
    clean_text: str
    blocks: list[DocumentBlock]
    multimodal_blocks: list[dict]
    tables: list[dict]
    features: dict[str, Any]


def clean_text(value: str) -> str:
    if not value:
        return ""

    replacements = {
        "\u00a0": " ",
        "\ufeff": "",
        "L\u00edquida": "Liquida",
        "Lquida": "Liquida",
        "milhes": "milhoes",
        "milh\u00f5es": "milhoes",
        "Ajustada": "Ajustada",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)

    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    value = re.sub(r"(?m)^[ \t]+|[ \t]+$", "", value)
    return value.strip()


def compact_text(value: str) -> str:
    return re.sub(r"\s+", " ", clean_text(value)).strip()


def is_heading(value: str) -> bool:
    text = compact_text(value)
    if not text:
        return False

    words = text.split()
    if len(words) > 14:
        return False

    if re.match(r"^\d+(\.\d+)*\s+.+", text):
        return True

    uppercase_letters = sum(1 for char in text if char.isupper())
    letters = sum(1 for char in text if char.isalpha())
    return letters > 0 and uppercase_letters / letters >= 0.7


def build_multimodal(block: DocumentBlock, position: int) -> dict:
    return {
        "block_type": block.block_type,
        "page_number": block.page_number,
        "content_text": block.clean_text or None,
        "metadata": {
            **block.metadata,
            "position": position,
            "section_title": block.section_title,
        },
    }


def parse_pdf(content: bytes) -> ParsedDocument:
    blocks: list[DocumentBlock] = []
    tables: list[dict] = []
    current_section: str | None = None

    with fitz.open(stream=content, filetype="pdf") as pdf:
        for page_index, page in enumerate(pdf, start=1):
            page_text = page.get_text("text").strip()
            for raw_block in page.get_text("blocks"):
                text = clean_text(raw_block[4])
                if not text:
                    continue

                if is_heading(text):
                    current_section = compact_text(text)

                blocks.append(DocumentBlock(
                    block_type="heading" if is_heading(text) else "text",
                    raw_text=raw_block[4],
                    clean_text=text,
                    page_number=page_index,
                    section_title=current_section,
                    metadata={
                        "source": "pymupdf",
                        "bbox": [raw_block[0], raw_block[1], raw_block[2], raw_block[3]],
                        "char_count": len(text),
                    },
                ))

            try:
                for table_index, table in enumerate(page.find_tables().tables, start=1):
                    rows = table.extract()
                    table_text = "\n".join(" | ".join(str(cell or "") for cell in row) for row in rows)
                    table_block = DocumentBlock(
                        block_type="table",
                        raw_text=table_text,
                        clean_text=clean_text(table_text),
                        page_number=page_index,
                        section_title=current_section,
                        metadata={"source": "pymupdf", "table_index": table_index, "row_count": len(rows)},
                    )
                    blocks.append(table_block)
                    tables.append({"page_number": page_index, "rows": rows, "section_title": current_section})
            except Exception as error:
                blocks.append(DocumentBlock(
                    block_type="table_error",
                    raw_text="",
                    clean_text="",
                    page_number=page_index,
                    section_title=current_section,
                    metadata={"source": "pymupdf", "error": str(error)},
                ))

            for image_index, _ in enumerate(page.get_images(full=True), start=1):
                blocks.append(DocumentBlock(
                    block_type="image",
                    raw_text="",
                    clean_text="",
                    page_number=page_index,
                    section_title=current_section,
                    metadata={
                        "source": "pymupdf",
                        "image_index": image_index,
                        "description": "Image detected. Vision interpretation is reserved for the provider-backed multimodal step.",
                    },
                ))

            if page_text and not any(block.page_number == page_index for block in blocks):
                blocks.append(DocumentBlock("text", page_text, clean_text(page_text), page_index, current_section))

    return finalize(blocks, tables, source="pdf")


def parse_docx(content: bytes) -> ParsedDocument:
    document = docx.Document(io.BytesIO(content))
    blocks: list[DocumentBlock] = []
    tables: list[dict] = []
    current_section: str | None = None

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        if is_heading(text):
            current_section = compact_text(text)
        blocks.append(DocumentBlock(
            block_type="heading" if is_heading(text) else "text",
            raw_text=paragraph.text,
            clean_text=text,
            section_title=current_section,
            metadata={"style": paragraph.style.name if paragraph.style else None},
        ))

    for table_index, table in enumerate(document.tables, start=1):
        rows = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
        table_text = "\n".join(" | ".join(row) for row in rows)
        blocks.append(DocumentBlock(
            block_type="table",
            raw_text=table_text,
            clean_text=clean_text(table_text),
            section_title=current_section,
            metadata={"source": "docx", "table_index": table_index, "row_count": len(rows)},
        ))
        tables.append({"rows": rows, "section_title": current_section})

    return finalize(blocks, tables, source="docx")


def parse_url(url: str) -> ParsedDocument:
    response = requests.get(url, timeout=10)
    response.raise_for_status()
    soup = BeautifulSoup(response.content, "html.parser")
    blocks: list[DocumentBlock] = []
    current_section: str | None = None

    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
        text = clean_text(element.get_text(separator=" ", strip=True))
        if not text:
            continue

        if element.name in {"h1", "h2", "h3", "h4"} or is_heading(text):
            current_section = compact_text(text)

        blocks.append(DocumentBlock(
            block_type="table" if element.name == "table" else ("heading" if element.name.startswith("h") else "text"),
            raw_text=text,
            clean_text=text,
            section_title=current_section,
            metadata={"source": "url", "tag": element.name, "url": url},
        ))

    return finalize(blocks, [], source="url")


def parse_text(content: bytes | str, source: str = "text") -> ParsedDocument:
    raw = content.decode("utf-8", errors="ignore") if isinstance(content, bytes) else content
    blocks: list[DocumentBlock] = []
    current_section: str | None = None

    for part in re.split(r"\n\s*\n+", raw):
        text = clean_text(part)
        if not text:
            continue
        if is_heading(text):
            current_section = compact_text(text)
        blocks.append(DocumentBlock(
            block_type="heading" if is_heading(text) else "text",
            raw_text=part,
            clean_text=text,
            section_title=current_section,
            metadata={"source": source},
        ))

    return finalize(blocks, [], source=source)


def finalize(blocks: list[DocumentBlock], tables: list[dict], source: str) -> ParsedDocument:
    visible_blocks = [block for block in blocks if block.clean_text]
    raw_text = "\n\n".join(block.raw_text for block in visible_blocks if block.raw_text)
    normalized = "\n\n".join(block.clean_text for block in visible_blocks)
    multimodal = [build_multimodal(block, index) for index, block in enumerate(blocks)]
    features = {
        "source": source,
        "block_count": len(blocks),
        "text_block_count": sum(1 for block in blocks if block.block_type == "text"),
        "heading_count": sum(1 for block in blocks if block.block_type == "heading"),
        "table_count": len(tables),
        "image_count": sum(1 for block in blocks if block.block_type == "image"),
        "char_count": len(normalized),
    }
    return ParsedDocument(
        raw_text=raw_text,
        clean_text=normalized,
        blocks=visible_blocks,
        multimodal_blocks=multimodal,
        tables=tables,
        features=features,
    )
